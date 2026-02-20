"""
Textract Service — Document Text Extraction
=============================================

Provides methods for extracting text from documents using both
Amazon Textract (for PDFs and images) and python-docx (for .docx files).

Note: Amazon Textract does NOT support .docx files directly.
This service handles the two-path approach:
1. .docx files → python-docx for direct text extraction
2. PDF/image files → Amazon Textract for OCR-based extraction

Usage:
    >>> from shared.textract_service import TextractService
    >>> textract = TextractService()
    >>> text = textract.convert_docx_to_text("path/to/file.docx")
    >>> result = textract.process_document_textract("bucket", "key.pdf")
"""

import os
import io
import logging
from typing import Optional

import boto3
from botocore.exceptions import ClientError

logger = logging.getLogger(__name__)


class TextractService:
    """
    Service class for document text extraction operations.

    Handles both direct .docx parsing (via python-docx) and
    Textract API calls (for PDF/image documents).

    Attributes:
        client: boto3 Textract client.
        bucket_name (str): Default S3 bucket for document storage.
    """

    def __init__(self, bucket_name: Optional[str] = None):
        """
        Initialize Textract service.

        Args:
            bucket_name: S3 bucket for document storage. Defaults to BUCKET_NAME env var.
        """
        self.client = boto3.client("textract")
        self.bucket_name = bucket_name or os.environ.get("BUCKET_NAME", "interview-iq-docs")
        logger.info("TextractService initialized with bucket: %s", self.bucket_name)

    def convert_docx_to_text(self, docx_path: str) -> str:
        """
        Extract text from a .docx file using python-docx.

        This is the primary method for .docx files since Textract doesn't
        support .docx format directly. Preserves paragraph structure.

        Args:
            docx_path: Local file path to the .docx document.

        Returns:
            str: Extracted text with paragraphs separated by newlines.

        Raises:
            FileNotFoundError: If the file doesn't exist.
            Exception: If python-docx cannot parse the file.
        """
        from docx import Document

        doc = Document(docx_path)
        paragraphs = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)

        result = "\n".join(paragraphs)
        logger.info("Extracted %d paragraphs (%d chars) from %s", len(paragraphs), len(result), docx_path)
        return result

    def convert_docx_bytes_to_text(self, docx_bytes: bytes) -> str:
        """
        Extract text from .docx file bytes (e.g., from S3 download).

        Args:
            docx_bytes: Raw bytes of a .docx file.

        Returns:
            str: Extracted text with paragraphs separated by newlines.
        """
        from docx import Document

        doc = Document(io.BytesIO(docx_bytes))
        paragraphs = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)

        result = "\n".join(paragraphs)
        logger.info("Extracted %d paragraphs (%d chars) from bytes", len(paragraphs), len(result))
        return result

    def process_document_textract(self, bucket: str, key: str) -> dict:
        """
        Process a document (PDF or image) using Amazon Textract.

        Calls the Textract DetectDocumentText API to perform OCR on
        the specified S3 object. Only works with PDF, PNG, JPEG, TIFF.

        Args:
            bucket: S3 bucket containing the document.
            key: S3 key of the document.

        Returns:
            dict: Extraction results with:
                - text (str): Full extracted text
                - blocks (list): Raw Textract Block objects
                - confidence (float): Average confidence score
                - page_count (int): Number of pages processed

        Raises:
            ClientError: If Textract API call fails.
        """
        try:
            response = self.client.detect_document_text(
                Document={"S3Object": {"Bucket": bucket, "Name": key}}
            )

            blocks = response.get("Blocks", [])
            lines = []
            confidences = []

            for block in blocks:
                if block["BlockType"] == "LINE":
                    lines.append(block["Text"])
                    confidences.append(block.get("Confidence", 0))

            text = "\n".join(lines)
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0

            result = {
                "text": text,
                "blocks": blocks,
                "confidence": round(avg_confidence, 2),
                "page_count": len(set(b.get("Page", 1) for b in blocks)),
                "line_count": len(lines),
            }

            logger.info(
                "Textract processed %s/%s: %d lines, %.1f%% confidence",
                bucket, key, len(lines), avg_confidence,
            )
            return result

        except ClientError as e:
            logger.error("Textract failed for %s/%s: %s", bucket, key, e)
            raise

    def upload_and_process(
        self,
        file_bytes: bytes,
        filename: str,
        session_id: str,
    ) -> dict:
        """
        Upload a document to S3 and extract its text.

        Routes to the appropriate extraction method based on file type:
        - .docx → python-docx (direct extraction)
        - .pdf, .png, .jpg, .tiff → Textract (OCR)

        Args:
            file_bytes: Raw file content as bytes.
            filename: Original filename (used to determine file type).
            session_id: Session ID for S3 organization.

        Returns:
            dict: Extraction results with:
                - text (str): Extracted text content
                - method (str): "python-docx" or "textract"
                - filename (str): Original filename
                - s3_key (str): Where the file was stored
        """
        # Upload to S3 first
        s3_key = f"uploads/{session_id}/{filename}"
        s3_client = boto3.client("s3")
        s3_client.put_object(Bucket=self.bucket_name, Key=s3_key, Body=file_bytes)
        logger.info("Uploaded %s to s3://%s/%s", filename, self.bucket_name, s3_key)

        # Route based on file extension
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

        if ext == "docx":
            # Direct extraction — faster and doesn't require Textract
            text = self.convert_docx_bytes_to_text(file_bytes)
            return {
                "text": text,
                "method": "python-docx",
                "filename": filename,
                "s3_key": s3_key,
            }
        elif ext in ("pdf", "png", "jpg", "jpeg", "tiff"):
            # Textract OCR extraction
            result = self.process_document_textract(self.bucket_name, s3_key)
            result["method"] = "textract"
            result["filename"] = filename
            result["s3_key"] = s3_key
            return result
        else:
            logger.warning("Unsupported file type: %s", ext)
            return {
                "text": "",
                "method": "unsupported",
                "filename": filename,
                "s3_key": s3_key,
                "error": f"Unsupported file type: .{ext}",
            }
